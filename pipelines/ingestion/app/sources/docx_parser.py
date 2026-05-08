"""DOCX parser for module user guides (text-first with selective OCR hooks)."""

from __future__ import annotations

import logging
import re
import zipfile
from datetime import datetime
import io
from pathlib import Path
from typing import List
from urllib.parse import quote

from docx import Document
from PIL import Image
import pytesseract

import sys

sys.path.insert(0, "/home/phungkien/ehc_ai_mvp")

from shared.py.clients.ollama_client import OllamaVision
from shared.py.models.faq import ModuleDocSection, stable_source_id
from app.utils.image_storage import ImageStorage, ImageExtractor

logger = logging.getLogger(__name__)

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _looks_heading(style_name: str) -> bool:
    style = (style_name or "").strip().lower()
    return style.startswith("heading") or style.startswith("tiêu đề")


def _module_name_from_file(file_path: Path) -> str:
    stem = file_path.stem
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem


def _should_ocr_image(file_name: str, image_bytes: bytes) -> bool:
    ext = Path(file_name).suffix.lower()
    if ext not in _IMAGE_EXTS:
        return False
    size = len(image_bytes)
    # Skip very small logos/icons and overly large binaries in POC path.
    return 15_000 <= size <= 3_000_000


def _ocr_with_tesseract(image_bytes: bytes) -> str:
    """CPU OCR fallback for environments without a running vision LLM endpoint."""
    try:
        img = Image.open(io.BytesIO(image_bytes))
        # Vietnamese + English works well for mixed UI labels.
        text = pytesseract.image_to_string(img, lang="vie+eng")
        return _clean_text(text)
    except Exception:
        return ""


class DocxParser:
    """Parse DOCX manuals into structured sections for ingestion."""

    def __init__(
        self,
        *,
        provider: str,
        vision_base_url: str,
        vision_model: str,
        enable_ocr: bool = True,
        max_ocr_images: int = 6,
        image_storage: ImageStorage = None,
    ):
        self.provider = (provider or "ollama").strip().lower()
        self.enable_ocr = enable_ocr
        self.max_ocr_images = max(0, int(max_ocr_images))
        self.image_storage = image_storage
        self._vision = None
        self._ocr_endpoint_ready = True
        if self.enable_ocr and self.max_ocr_images > 0:
            self._vision = OllamaVision(
                base_url=vision_base_url,
                model=vision_model,
                timeout=90,
                provider=self.provider,
            )

    def parse_file(self, file_path: str) -> List[ModuleDocSection]:
        docx_path = Path(file_path)
        if not docx_path.exists() or docx_path.suffix.lower() != ".docx":
            logger.warning("Skip non-docx path: %s", file_path)
            return []

        stat = docx_path.stat()
        created_at = datetime.fromtimestamp(stat.st_ctime)
        updated_at = datetime.fromtimestamp(stat.st_mtime)
        module_name = _module_name_from_file(docx_path)

        # Extract images with OCR in a single pass — avoids re-OCR in _extract_ocr_sections
        image_ids = []
        if self.image_storage:
            ocr_callback = self._ocr_image if self.enable_ocr else None
            extractor = ImageExtractor(self.image_storage)
            image_ids, _ = extractor.extract_from_docx(
                str(docx_path),
                ocr_callback=ocr_callback,
                max_images=self.max_ocr_images,
            )
            logger.info("Extracted %d images from %s", len(image_ids), docx_path.name)

        # Text sections carry no images — mapping images to sections requires DOCX XML parsing
        sections = self._extract_text_sections(docx_path, module_name, created_at, updated_at)
        if image_ids and self.image_storage:
            # Build OCR section from already-stored alt_texts — no second vision call
            ocr_section = self._build_ocr_section_from_storage(
                docx_path, module_name, created_at, updated_at, image_ids
            )
            if ocr_section:
                sections.append(ocr_section)
        else:
            ocr_sections = self._extract_ocr_sections(docx_path, module_name, created_at, updated_at)
            sections.extend(ocr_sections)

        logger.info(
            "Parsed DOCX %s into %d sections (images=%d)",
            docx_path.name,
            len(sections),
            len(image_ids),
        )
        return sections

    def parse_directory(self, docx_dir: str) -> List[ModuleDocSection]:
        root = Path(docx_dir)
        if not root.exists():
            logger.warning("DOCX directory not found: %s", docx_dir)
            return []

        all_sections: List[ModuleDocSection] = []
        for path in sorted(root.glob("*.docx")):
            all_sections.extend(self.parse_file(str(path)))
        return all_sections

    def _ocr_image(self, image_bytes: bytes) -> str:
        """OCR callback for ImageExtractor."""
        if not self._vision or not self._ocr_endpoint_ready:
            return _ocr_with_tesseract(image_bytes)
        
        try:
            ocr_text = self._vision.extract_text_from_image(
                image_bytes=image_bytes,
                prompt=(
                    "Trich xuat ngan gon text UI quan trong tu anh huong dan phan mem. "
                    "Uu tien: ten menu, ten nut, ten truong nhap, thong bao loi neu co. "
                    "Khong suy dien them."
                ),
            )
            return _clean_text(ocr_text)
        except Exception as ocr_err:
            err_msg = str(ocr_err)
            if "404" in err_msg and "/v1/chat/completions" in err_msg:
                self._ocr_endpoint_ready = False
                logger.warning(
                    "DOCX OCR LLM disabled: vision endpoint does not support chat completions; switching to Tesseract fallback"
                )
            else:
                logger.debug("OCR error: %s", ocr_err)
            
            return _ocr_with_tesseract(image_bytes)

    def _extract_text_sections(
        self,
        docx_path: Path,
        module_name: str,
        created_at: datetime,
        updated_at: datetime,
    ) -> List[ModuleDocSection]:
        document = Document(str(docx_path))

        sections: List[ModuleDocSection] = []
        section_title = "Tổng quan"
        section_lines: List[str] = []
        section_idx = 0

        def flush_section() -> None:
            nonlocal section_idx, section_lines, section_title
            body = "\n".join([line for line in section_lines if line]).strip()
            if not body:
                section_lines = []
                return
            section_idx += 1
            source_id = stable_source_id(
                module_name,
                section_title,
                str(section_idx),
                prefix="docx",
            )
            section_url = f"docx://{quote(docx_path.name)}#section={quote(section_title)}"
            sections.append(
                ModuleDocSection(
                    source_id=source_id,
                    module_name=module_name,
                    section_title=section_title,
                    content=body,
                    source_file=docx_path.name,
                    source_url=section_url,
                    created_at=created_at,
                    updated_at=updated_at,
                    image_ids=[],
                    metadata={
                        "source_type": "module_doc",
                        "has_images": False,
                        "image_count": 0,
                        "section_index": section_idx,
                    },
                )
            )
            section_lines = []

        for para in document.paragraphs:
            text = _clean_text(para.text)
            if not text:
                continue

            if _looks_heading(getattr(para.style, "name", "")):
                flush_section()
                section_title = text
                continue

            section_lines.append(text)

        # Preserve table text (flattened) as extra sections for retrieval.
        table_idx = 0
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [_clean_text(cell.text) for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            table_idx += 1
            flush_section()
            section_title = f"Bảng {table_idx}"
            section_lines.extend(rows)

        flush_section()
        return sections

    def _build_ocr_section_from_storage(
        self,
        docx_path: Path,
        module_name: str,
        created_at: datetime,
        updated_at: datetime,
        image_ids: List[str],
    ) -> "ModuleDocSection | None":
        """Build OCR section from already-stored image alt_texts (no second vision call)."""
        snippets: List[str] = []
        for img_id in image_ids:
            meta = self.image_storage.get_image_metadata(img_id)
            alt_text = (meta.get("alt_text") or "").strip()
            orig_name = meta.get("original_filename", img_id)
            if len(alt_text) >= 20:
                snippets.append(f"{orig_name}: {alt_text[:700]}")

        if not snippets:
            return None

        section_title = "OCR hình minh họa"
        content = "\n".join([f"- {s}" for s in snippets])
        source_id = stable_source_id(module_name, section_title, prefix="docx")
        section_url = f"docx://{quote(docx_path.name)}#section={quote(section_title)}"

        return ModuleDocSection(
            source_id=source_id,
            module_name=module_name,
            section_title=section_title,
            content=content,
            source_file=docx_path.name,
            source_url=section_url,
            created_at=created_at,
            updated_at=updated_at,
            image_ids=image_ids.copy(),
            metadata={
                "source_type": "module_doc",
                "has_images": True,
                "image_ocr_coverage": len(snippets),
            },
        )

    def _extract_ocr_sections(
        self,
        docx_path: Path,
        module_name: str,
        created_at: datetime,
        updated_at: datetime,
    ) -> List[ModuleDocSection]:
        if not self.enable_ocr:
            return []

        snippets: List[str] = []
        try:
            with zipfile.ZipFile(str(docx_path), "r") as zf:
                media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
                for media_name in media_files:
                    if len(snippets) >= self.max_ocr_images:
                        break
                    image_bytes = zf.read(media_name)
                    if not _should_ocr_image(media_name, image_bytes):
                        continue
                    cleaned = ""
                    try:
                        if self._vision and self._ocr_endpoint_ready:
                            ocr_text = self._vision.extract_text_from_image(
                                image_bytes=image_bytes,
                                prompt=(
                                    "Trich xuat ngan gon text UI quan trong tu anh huong dan phan mem. "
                                    "Uu tien: ten menu, ten nut, ten truong nhap, thong bao loi neu co. "
                                    "Khong suy dien them."
                                ),
                            )
                            cleaned = _clean_text(ocr_text)
                    except Exception as ocr_err:
                        err_msg = str(ocr_err)
                        if "404" in err_msg and "/v1/chat/completions" in err_msg:
                            # Current endpoint is embedding-only. Disable OCR for the
                            # rest of this run to avoid repeated slow failures.
                            self._ocr_endpoint_ready = False
                            logger.warning(
                                "DOCX OCR LLM disabled for this run: vision endpoint does not support chat completions; switching to Tesseract fallback"
                            )
                        else:
                            logger.debug("OCR skip image %s: %s", media_name, ocr_err)

                    if len(cleaned) < 20:
                        cleaned = _ocr_with_tesseract(image_bytes)

                    if len(cleaned) >= 20:
                        snippets.append(f"{Path(media_name).name}: {cleaned[:700]}")
        except Exception as exc:
            logger.warning("Failed to read DOCX media for %s: %s", docx_path.name, exc)
            return []

        if not snippets:
            return []

        section_title = "OCR hình minh họa"
        content = "\n".join([f"- {snippet}" for snippet in snippets])
        source_id = stable_source_id(module_name, section_title, prefix="docx")
        section_url = f"docx://{quote(docx_path.name)}#section={quote(section_title)}"

        return [
            ModuleDocSection(
                source_id=source_id,
                module_name=module_name,
                section_title=section_title,
                content=content,
                source_file=docx_path.name,
                source_url=section_url,
                created_at=created_at,
                updated_at=updated_at,
                metadata={
                    "source_type": "module_doc",
                    "has_images": True,
                    "image_ocr_coverage": len(snippets),
                },
            )
        ]
